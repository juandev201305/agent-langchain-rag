from docx import Document
from io import BytesIO
from schemas.generated_document import GeneratedDocument
from services.document_store import save_document
import uuid
from datetime import datetime

# Genera un archivo DOCX en memoria a partir del modelo estructurado
def generate_docx(doc_model: DocumentModel) -> GeneratedDocument:
    doc = Document()

    # Titulo principal
    doc.add_heading(doc_model.title, level=1)

    # Secciones
    for section in doc_model.sections:
        doc.add_heading(section.title, level=2)

        for paragraph in section.paragraphs:
            doc.add_paragraph(paragraph.text)
    
    buffer = BytesIO()
    # Guardar documento generado en memoria
    doc.save(buffer)
    buffer.seek(0)

    document = GeneratedDocument(
        id=str(uuid.uuid4()),
        content=buffer,
        filename="documento.docx",
        created_at=datetime.utcnow(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    print(document)
    save_document(document)

    return document.id
